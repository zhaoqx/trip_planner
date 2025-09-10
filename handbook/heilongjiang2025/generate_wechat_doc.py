import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import pillow_heif
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE_DIR, 'blogs', '哈尔滨1.md')
IMG_DIR = os.path.join(BASE_DIR, 'blogs', 'pic', '哈尔滨1')
OUTPUT_PATH = os.path.join(BASE_DIR, 'blogs', '哈尔滨1_公众号排版版.docx')

# Mapping placeholder keywords to actual file names (chosen images)
PLACEHOLDER_IMAGE_MAP = {
    '酒店江景': '酒店江景.jpg',  # 指定用户提供的文件
    '老厨家的菜': '老厨家的菜.jpeg',
    '索菲亚教堂外景': '索菲亚教堂外景.jpeg',
    '松花江畔夕阳': '松花江畔夕阳.jpeg',
    '铁路桥': '铁路桥.jpeg',
    '张包铺': '张包铺.jpg',
}

# Additional context-driven images to insert (order sensitive)
EXTRA_INSERTIONS = []  # 需求：去除“旅途剪影”等额外附加图片


def convert_if_heic(path):
    if not path.lower().endswith('.heic'):
        return path
    target = os.path.splitext(path)[0] + '.jpeg'
    if os.path.exists(target):
        return target
    try:
        heif_file = pillow_heif.read_heif(path)
        image = Image.frombytes(heif_file.mode, heif_file.size, heif_file.data, 'raw', heif_file.mode, heif_file.stride)
        image.save(target, 'JPEG')
        return target
    except Exception as e:
        print(f"Failed to convert {path}: {e}")
        return path


def setup_styles(doc: Document):
    """Adjust styles to more typical WeChat-friendly Chinese fonts.
    Title: 黑体  Body: 宋体  Caption: 仿宋 / 等线 fallback"""
    styles = doc.styles
    from docx.enum.style import WD_STYLE_TYPE

    # Title style
    try:
        title_style = styles['WeChatTitle']
    except KeyError:
        title_style = styles.add_style('WeChatTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = '黑体'
    try:
        title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    except Exception:
        pass
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(10)

    # Body style (Normal)
    body = styles['Normal']
    body.font.name = '宋体'
    try:
        body._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass
    body.font.size = Pt(13)  # 增加字号
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_after = Pt(10)  # 调整为稍大段后

    # Caption style
    try:
        cap = styles['WeChatCaption']
    except KeyError:
        cap = styles.add_style('WeChatCaption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = '仿宋'
    try:
        cap._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    except Exception:
        pass
    cap.font.size = Pt(9)
    cap.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)


def add_image(doc: Document, img_filename: str, caption: str | None = None, width_in=5.6):
    path = os.path.join(IMG_DIR, img_filename)
    if not os.path.exists(path):
        print(f"Image missing: {img_filename}")
        return
    path = convert_if_heic(path)
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
        if caption:
            doc.add_paragraph(caption, style='WeChatCaption')
    except Exception as e:
        print(f"Failed to add image {img_filename}: {e}")


def build_document():
    if not os.path.exists(MD_PATH):
        print('Markdown not found')
        return
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    setup_styles(doc)
    # 按需求：不添加标题

    lines = content.splitlines()

    # Track which placeholders handled
    used_from_map = set()

    for line in lines:
        stripped = line.strip()
        # 跳过原始一级标题行
        if stripped.startswith('# '):
            continue
        if stripped.startswith('[图片：') and stripped.endswith(']'):
            key = stripped.removeprefix('[图片：').removesuffix(']')
            # Try placeholder map first
            img_name = None
            for k,v in PLACEHOLDER_IMAGE_MAP.items():
                if k in key:
                    img_name = v
                    used_from_map.add(k)
                    break
            if img_name:
                add_image(doc, img_name, key)
            else:
                # fallback: try fuzzy match
                fallback = None
                for cand in os.listdir(IMG_DIR):
                    if any(seg in cand for seg in key.split()):
                        fallback = cand
                        break
                if fallback:
                    add_image(doc, fallback, key)
            continue
        if stripped == '':
            # 按需求：不保留空白段，直接跳过
            continue
        doc.add_paragraph(stripped)

    # 已按需求移除“旅途剪影”附加段落与额外图片插入

    doc.save(OUTPUT_PATH)
    print('Generated:', OUTPUT_PATH)

if __name__ == '__main__':
    build_document()
