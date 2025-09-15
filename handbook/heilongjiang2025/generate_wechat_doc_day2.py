import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import pillow_heif
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(BASE_DIR, 'blogs', '哈尔滨2.md')
IMG_DIR = os.path.join(BASE_DIR, 'blogs', 'pic', '哈尔滨2')
OUTPUT_PATH = os.path.join(BASE_DIR, 'blogs', '哈尔滨2_公众号排版版.docx')

# Mapping placeholder keywords to actual file names (chosen images)
PLACEHOLDER_IMAGE_MAP = {
    '东北虎林园喂食': '东北虎园喂食.jpg',
}

# regex: [图片:关键词] or [图片：关键词] optionally followed by (path) or （path）
IMG_MARKER_RE = re.compile(r"^\[图片[:：]\s*(.+?)\](?:[\(\（]\s*(.+?)\s*[\)\）])?$")


def convert_if_heic(path: str) -> str:
    if not path.lower().endswith('.heic'):
        return path
    target = os.path.splitext(path)[0] + '.jpeg'
    if os.path.exists(target):
        return target
    try:
        heif_file = pillow_heif.read_heif(path)
        image = Image.frombytes(heif_file.mode, heif_file.size, heif_file.data, 'raw', heif_file.mode, heif_file.stride)
        image.save(target, 'JPEG')
        return target
    except Exception as e:
        print(f"Failed to convert {path}: {e}")
        return path


def setup_styles(doc: Document):
    """Adjust styles to WeChat-friendly Chinese fonts.
    Body: 宋体 13pt, 1.5x line spacing; Caption: 仿宋 9pt, centered."""
    styles = doc.styles
    from docx.enum.style import WD_STYLE_TYPE

    # Body style (Normal)
    body = styles['Normal']
    body.font.name = '宋体'
    try:
        body._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass
    body.font.size = Pt(13)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_after = Pt(10)

    # Caption style
    try:
        cap = styles['WeChatCaption']
    except KeyError:
        cap = styles.add_style('WeChatCaption', WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = '仿宋'
    try:
        cap._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    except Exception:
        pass
    cap.font.size = Pt(9)
    cap.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)


def add_image(doc: Document, img_ref: str, caption: str | None = None, width_in=5.6):
    # If img_ref includes a path separator, treat it as a direct file name in this project's tree and use basename under IMG_DIR
    img_name = os.path.basename(img_ref)
    candidate_paths = [
        os.path.join(IMG_DIR, img_name),
        os.path.join(BASE_DIR, 'blogs', 'pic', '哈尔滨2', img_name),
        os.path.join(BASE_DIR, img_ref) if not os.path.isabs(img_ref) else img_ref,
    ]
    path = None
    for p in candidate_paths:
        if os.path.exists(p):
            path = p
            break
    if path is None:
        # Last resort: search by name in IMG_DIR
        for cand in os.listdir(IMG_DIR) if os.path.isdir(IMG_DIR) else []:
            if img_name in cand:
                path = os.path.join(IMG_DIR, cand)
                break
    if path is None:
        print(f"Image missing: {img_ref}")
        return

    path = convert_if_heic(path)
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
        if caption:
            doc.add_paragraph(caption, style='WeChatCaption')
    except Exception as e:
        print(f"Failed to add image {img_ref}: {e}")


def parse_image_marker(line: str):
    m = IMG_MARKER_RE.match(line.strip())
    if not m:
        return None
    key = m.group(1).strip()
    raw_path = m.group(2).strip() if m.group(2) else None
    return key, raw_path


def refine_text(text: str) -> str:
    # Targeted light edits without changing structure/meaning
    fixes = [
        (r"做轮渡", "坐轮渡"),
        (r"东北虎园(?!林园)", "东北虎林园"),
        (r"\s+", " "),  # collapse multiple spaces
    ]
    out = text
    for pat, rep in fixes:
        out = re.sub(pat, rep, out)
    return out.strip()


def build_document():
    if not os.path.exists(MD_PATH):
        print('Markdown not found:', MD_PATH)
        return
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    setup_styles(doc)

    lines = content.splitlines()

    for line in lines:
        stripped = line.strip()
        # Skip original H1 title lines
        if stripped.startswith('# '):
            continue
        # Handle image markers
        parsed = parse_image_marker(stripped)
        if parsed:
            key, raw_path = parsed
            img_name = None
            # explicit path in marker takes precedence
            if raw_path:
                img_name = os.path.basename(raw_path)
            else:
                # try placeholder map
                if key in PLACEHOLDER_IMAGE_MAP:
                    img_name = PLACEHOLDER_IMAGE_MAP[key]
                else:
                    # fuzzy search within IMG_DIR
                    for cand in os.listdir(IMG_DIR) if os.path.isdir(IMG_DIR) else []:
                        if all(seg in cand for seg in key.split()):
                            img_name = cand
                            break
            if img_name:
                add_image(doc, img_name, key)
            else:
                print(f"No image matched for placeholder: {key}")
            continue
        # skip empty lines (compact layout)
        if stripped == '':
            continue
        # light copy-edit
        refined = refine_text(stripped)
        doc.add_paragraph(refined)

    doc.save(OUTPUT_PATH)
    print('Generated:', OUTPUT_PATH)


if __name__ == '__main__':
    build_document()
